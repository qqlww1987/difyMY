import logging
import os
import tqdm
from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.models.document import Document

logger = logging.getLogger(__name__)


class UnstructuredWordExtractor(BaseExtractor):
    """Loader that uses unstructured to load word documents.
    """

    def __init__(
            self,
            file_path: str,
            api_url: str,
    ):
        """Initialize with file path."""
        self._file_path = file_path
        self._api_url = api_url
    def extract(self) -> list[Document]:
        from unstructured.__version__ import __version__ as __unstructured_version__
        from unstructured.file_utils.filetype import FileType, detect_filetype
        def doc2text(filepath):
            from docx.table import _Cell, Table
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx import Document, ImagePart
            from PIL import Image
            from io import BytesIO
            import numpy as np
            # from rapidocr_onnxruntime import RapidOCR
            # ocr = RapidOCR()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(total=len(doc.paragraphs)+len(doc.tables),
                               desc="RapidOCRDocLoader block index: 0")
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    # images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    # for image in images:
                    #     for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                    #         part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                    #         if isinstance(part, ImagePart):
                    #             image = Image.open(BytesIO(part._blob))
                    #             result, _ = ocr(np.array(image))
                    #             if result:
                    #                 ocr_result = [line[1] for line in result]
                    #                 resp += "\n".join(ocr_result)
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                b_unit.update(1)
            return resp
    
        unstructured_version = tuple(
            [int(x) for x in __unstructured_version__.split(".")]
        )
        # check the file extension
        try:
            import magic  # noqa: F401

            is_doc = detect_filetype(self._file_path) == FileType.DOC
        except ImportError:
            _, extension = os.path.splitext(str(self._file_path))
            is_doc = extension == ".doc"

        if is_doc and unstructured_version < (0, 4, 11):
            raise ValueError(
                f"You are on unstructured version {__unstructured_version__}. "
                "Partitioning .doc files is only supported in unstructured>=0.4.11. "
                "Please upgrade the unstructured package and try again."
            )
        resultString= doc2text(self._file_path)
        print(resultString)
        if is_doc:
            from unstructured.partition.doc import partition_doc

            elements = partition_doc(text=resultString, **self.unstructured_kwargs)
        else:
            from unstructured.partition.docx import partition_docx

            elements = partition_docx(text=resultString, **self.unstructured_kwargs)

        from unstructured.chunking.title import chunk_by_title
        chunks = chunk_by_title(elements, max_characters=2000, combine_text_under_n_chars=2000)
        documents = []
        for chunk in chunks:
            text = chunk.text.strip()
            documents.append(Document(page_content=text))
        return documents
   # def extract(self) -> list[Document]:
    #     from unstructured.__version__ import __version__ as __unstructured_version__
    #     from unstructured.file_utils.filetype import FileType, detect_filetype

    #     unstructured_version = tuple(
    #         [int(x) for x in __unstructured_version__.split(".")]
    #     )
    #     # check the file extension
    #     try:
    #         import magic  # noqa: F401

    #         is_doc = detect_filetype(self._file_path) == FileType.DOC
    #     except ImportError:
    #         _, extension = os.path.splitext(str(self._file_path))
    #         is_doc = extension == ".doc"

    #     if is_doc and unstructured_version < (0, 4, 11):
    #         raise ValueError(
    #             f"You are on unstructured version {__unstructured_version__}. "
    #             "Partitioning .doc files is only supported in unstructured>=0.4.11. "
    #             "Please upgrade the unstructured package and try again."
    #         )

    #     if is_doc:
    #         from unstructured.partition.doc import partition_doc

    #         elements = partition_doc(filename=self._file_path)
    #     else:
    #         from unstructured.partition.docx import partition_docx

    #         elements = partition_docx(filename=self._file_path)

    #     from unstructured.chunking.title import chunk_by_title
    #     chunks = chunk_by_title(elements, max_characters=2000, combine_text_under_n_chars=2000)
    #     documents = []
    #     for chunk in chunks:
    #         text = chunk.text.strip()
    #         documents.append(Document(page_content=text))
    #     return documents
