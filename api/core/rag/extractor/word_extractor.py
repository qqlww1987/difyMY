"""Abstract interface for document loader implementations."""
import os
import tempfile
import tqdm
import pandas as pd
from urllib.parse import urlparse

import requests

from core.rag.extractor.extractor_base import BaseExtractor
from core.rag.models.document import Document


class WordExtractor(BaseExtractor):
    """Load docx files.


    Args:
        file_path: Path to the file to load.
    """

    def __init__(self, file_path: str):
        """Initialize with file path."""
        self.file_path = file_path
        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

        # If the file is a web path, download it to a temporary file, and use that
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            r = requests.get(self.file_path)

            if r.status_code != 200:
                raise ValueError(
                    f"Check the url of your file; returned status code {r.status_code}"
                )

            self.web_path = self.file_path
            self.temp_file = tempfile.NamedTemporaryFile()
            self.temp_file.write(r.content)
            self.file_path = self.temp_file.name
        elif not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file or url")

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """Load given path as single page."""
        from docx import Document as docx_Document
        def doc2text(doc:docx_Document):
            from docx.table import _Cell, Table
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx import Document, ImagePart
            from PIL import Image
            from io import BytesIO
            import numpy as np
            # from rapidocr_onnxruntime import RapidOCR
            # ocr = RapidOCR()
            # resp = ""
            doc_texts=[]
            # def get_table_dataframe(table):
            #     date = []
            #     keys = None
            #     for i, row in enumerate(table.rows):
            #         # 获取表格一行的数据
            #         text = (cell.text for cell in row.cells)  # text为generator生成器类型
            #         # 判断是否是表头
            #         if i == 0:
            #             keys = tuple(text)
            #             continue
            #         date.append(dict(zip(keys, text)))  # zip方法，按列打包为元组的列表。再转换为字典
            #     df = pd.DataFrame(date)  # pd依赖的DataFrame方法将字典数据转换成列表集
            #     return df
            def iter_block_items(parent):
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(total=len(doc.paragraphs)+len(doc.tables),
                               desc="RapidOCRDocLoader block index: 0")
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    doc_texts.append(block.text.strip()+ "\n")
                    # images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    # for image in images:
                    #     for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                    #         part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                    #         if isinstance(part, ImagePart):
                    #             image = Image.open(BytesIO(part._blob))
                    #             result, _ = ocr(np.array(image))
                    #             if result:
                    #                 ocr_result = [line[1] for line in result]
                    #                 resp += "\n".join(ocr_result)
                if isinstance(block, Table):
       
                    resp=""
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                # 换行符格式化
                                resp += paragraph.text.strip()
                    doc_texts.append(resp)
                b_unit.update(1)
            return doc_texts
    
        document = docx_Document(self.file_path)
        doc_texts=doc2text(document)
        content = '\n'.join(doc_texts)
        # doc_texts = [paragraph.text for paragraph in document.paragraphs]
        # doc_texts.append(doc_Tabbele)
        # content = '\n'.join(doc_Tabbele)
        # print(content)# content = '\n'.join(doc_texts)
        return [Document(
            page_content=content,
            metadata={"source": self.file_path},
        )]

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)
